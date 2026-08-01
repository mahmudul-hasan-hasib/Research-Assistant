"""Document loaders (Part 6.1/6.2 — Loader → Preprocessor → Chunker).

Each source type is a distinct adapter behind the ``DocumentLoader`` port (OCP:
new file types add loaders, no core changes). PDF parsing delegates to LangChain's
``PyMuPDFLoader`` (PyMuPDF), DOCX parsing uses python-docx, plain text decodes
directly. Loaders accept raw bytes so they are storage-agnostic — the object store
hands them bytes (Part 4.2 storage port).
"""

from __future__ import annotations

import io
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Protocol

from app.core.exceptions import InvalidFileTypeError
from app.modules.uploads.validation import DOCX_TYPE, TEXT_TYPES, normalize_content_type

PARSER_BY_TYPE: dict[str, str] = {
    "application/pdf": "pdf",
    DOCX_TYPE: "docx",
    "text/plain": "text",
}


@dataclass
class LoadedDocument:
    """One parsing unit: a page for PDFs, a heading section for DOCX, the whole
    text for plain files."""

    page_content: str
    source: str
    page: int | None = None
    heading: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentLoader(Protocol):
    parser: str

    def load(self, data: bytes) -> list[LoadedDocument]: ...


class TextDocumentLoader:
    parser = "text"

    def __init__(self, *, encoding: str = "utf-8") -> None:
        self._encoding = encoding

    def load(self, data: bytes) -> list[LoadedDocument]:
        text = data.decode(self._encoding, errors="replace")
        return [LoadedDocument(page_content=text, source="text")]


class PDFDocumentLoader:
    """Parses a PDF into one ``LoadedDocument`` per page via LangChain/PyMuPDF.

    The loader preserves the page number in ``LoadedDocument.page`` so downstream
    chunks carry ``page`` metadata for citations (Part 6.2)."""

    parser = "pdf"

    def load(self, data: bytes) -> list[LoadedDocument]:
        from langchain_community.document_loaders import PyMuPDFLoader

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as handle:
            handle.write(data)
            tmp_path = Path(handle.name)
        try:
            documents = PyMuPDFLoader(str(tmp_path)).load()
        finally:
            tmp_path.unlink(missing_ok=True)

        loaded: list[LoadedDocument] = []
        for document in documents:
            page = document.metadata.get("page")
            loaded.append(
                LoadedDocument(
                    page_content=document.page_content or "",
                    source="pdf",
                    page=int(page) if page is not None else None,
                    metadata=dict(document.metadata),
                )
            )
        return loaded


class DOCXDocumentLoader:
    """Parses a DOCX into heading-delimited sections via python-docx.

    The upload allow-list already guarantees the bytes are a ZIP container
    (``docx`` magic check, Part 11); python-docx is the OOXML parser here."""

    parser = "docx"

    def load(self, data: bytes) -> list[LoadedDocument]:
        from docx import Document as DocxDocument

        document = DocxDocument(io.BytesIO(data))
        sections: list[tuple[str | None, list[str]]] = []
        current_heading: str | None = None
        current_lines: list[str] = []

        def flush() -> None:
            sections.append((current_heading, current_lines))

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            if style_name.startswith("heading"):
                flush()
                current_heading = text
                current_lines = []
            else:
                current_lines.append(text)
        flush()

        loaded: list[LoadedDocument] = []
        for heading, lines in sections:
            content = "\n".join(lines).strip()
            if content:
                loaded.append(
                    LoadedDocument(page_content=content, source="docx", heading=heading)
                )
        return loaded


def build_loader(content_type: str) -> DocumentLoader:
    """Return the loader for a validated content type (fail closed, P7)."""
    normalized = normalize_content_type(content_type)
    if normalized == "application/pdf":
        return PDFDocumentLoader()
    if normalized == DOCX_TYPE:
        return DOCXDocumentLoader()
    if normalized in TEXT_TYPES:
        return TextDocumentLoader()
    raise InvalidFileTypeError(detail=f"No loader for content type: {content_type}")
