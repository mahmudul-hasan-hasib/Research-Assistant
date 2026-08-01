"""RAG module tests: loaders, splitters, chunk manager, embeddings, vector store,
BM25, retrieval service, citations, and the ingestion→retrieval API flow."""

from __future__ import annotations

import uuid

import pytest
from fastapi.testclient import TestClient

from app.core.config import Settings
from app.core.exceptions import ConfigurationError
from app.main import create_app
from app.modules.llm.fake import FakeLLMProvider
from app.modules.rag.chunking import ChunkManager
from app.modules.rag.citations import CitationGenerator, RetrievalHit
from app.modules.rag.embeddings import HashingEmbedder, build_embedder
from app.modules.rag.loaders import (
    DOCXDocumentLoader,
    PDFDocumentLoader,
    TextDocumentLoader,
    build_loader,
)
from app.modules.rag.models import Document, DocumentStatus
from app.modules.rag.retrieval import (
    BM25,
    LLMQueryRewriter,
    RetrievalService,
)
from app.modules.rag.splitters import (
    LangChainTextSplitter,
    LlamaIndexTextSplitter,
    build_splitter,
    estimate_tokens,
)
from app.modules.rag.vectorstore import (
    FAISSVectorStore,
    InMemoryVectorStore,
    build_vector_store,
)
from app.shared.base import Base

CREDENTIALS = {"email": "user@example.com", "password": "correct-horse-battery", "display_name": "Ada"}

QUANTUM_TEXT = (
    "Quantum entanglement is a physical phenomenon where two or more particles "
    "become correlated such that the state of one cannot be described "
    "independently of the others. This correlation persists even when the "
    "particles are separated by large distances. Albert Einstein famously "
    "described entanglement as spooky action at a distance, while modern "
    "experiments continue to probe its role in quantum information science."
)


@pytest.fixture
def rag_settings(tmp_path) -> Settings:
    return Settings(
        app_name="insight-test",
        app_env="test",
        debug=False,
        log_level="CRITICAL",
        cors_origins=["http://localhost:3000"],
        database_url=f"sqlite:///{tmp_path / 'rag.db'}",
        jwt_secret_key="test-secret-with-at-least-32-characters",
        storage_backend="local",
        storage_local_root=str(tmp_path / "storage"),
        embedding_provider="memory",
        embedding_dimensions=32,
        vector_store_backend="memory",
        rag_splitter="langchain",
    )


@pytest.fixture
def client(rag_settings: Settings) -> TestClient:
    app = create_app(rag_settings)
    assert app.state.container.engine is not None
    Base.metadata.create_all(app.state.container.engine)
    with TestClient(app, raise_server_exceptions=False) as test_client:
        yield test_client


def _register(client: TestClient, **overrides) -> dict:
    response = client.post("/api/v1/auth/register", json={**CREDENTIALS, **overrides})
    assert response.status_code == 201, response.text
    return response.json()


def _ready_upload(client: TestClient, token: str, *, filename: str, content_type: str, data: bytes) -> dict:
    result = client.post(
        "/api/v1/uploads/presign",
        json={"filename": filename, "content_type": content_type, "size_bytes": len(data)},
        headers={"Authorization": f"Bearer {token}"},
    ).json()
    client.app.state.container.storage.put_bytes(result["storage_key"], data, content_type=content_type)
    response = client.post(
        f"/api/v1/uploads/{result['upload_id']}/complete",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 200, response.text
    return result


def _ingest(client: TestClient, token: str, upload_id: str) -> dict:
    response = client.post(
        "/api/v1/rag/documents",
        json={"upload_id": upload_id},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 201, response.text
    return response.json()


def _retrieve(client: TestClient, token: str, query: str, top_k: int = 5) -> dict:
    response = client.post(
        "/api/v1/rag/retrieve",
        json={"query": query, "top_k": top_k},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 200, response.text
    return response.json()


# --- loaders ----------------------------------------------------------------


def test_text_loader_decodes_utf8() -> None:
    loader = TextDocumentLoader()
    loaded = loader.load("héllo insight\nsecond line".encode())
    assert len(loaded) == 1
    assert loaded[0].page_content == "héllo insight\nsecond line"


def test_pdf_loader_keeps_page_numbers(tmp_path) -> None:
    import fitz

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "First page content for RAG.")
    page2 = document.new_page()
    page2.insert_text((72, 72), "Second page content for RAG.")
    data = document.tobytes()
    document.close()

    loaded = PDFDocumentLoader().load(data)
    assert len(loaded) == 2
    assert loaded[0].page == 0
    assert loaded[1].page == 1
    assert "First page content" in loaded[0].page_content


def test_docx_loader_splits_heading_sections() -> None:
    import io

    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_heading("Introduction", level=1)
    document.add_paragraph("Intro paragraph text.")
    document.add_heading("Methods", level=1)
    document.add_paragraph("Methods paragraph text.")
    buffer = io.BytesIO()
    document.save(buffer)

    loaded = DOCXDocumentLoader().load(buffer.getvalue())
    assert [unit.heading for unit in loaded] == ["Introduction", "Methods"]
    assert "Intro paragraph" in loaded[0].page_content


def test_build_loader_dispatches_and_fails_closed() -> None:
    assert build_loader("text/plain").parser == "text"
    assert build_loader("application/pdf").parser == "pdf"
    assert build_loader("application/vnd.openxmlformats-officedocument.wordprocessingml.document").parser == "docx"
    with pytest.raises(Exception) as exc:
        build_loader("application/x-msdownload")
    assert exc.value.status_code == 422


# --- splitters ---------------------------------------------------------------


def test_estimate_tokens() -> None:
    assert estimate_tokens("") == 1
    assert estimate_tokens("abcd") == 1
    assert estimate_tokens("abcdefgh") == 2


def test_langchain_splitter_respects_chunk_size() -> None:
    splitter = LangChainTextSplitter()
    pieces = splitter.split("word " * 100, chunk_size=50, chunk_overlap=5)
    assert len(pieces) > 1
    assert all(len(piece) <= 50 for piece in pieces)


def test_llamaindex_splitter_works_without_nltk() -> None:
    splitter = LlamaIndexTextSplitter()
    text = ". ".join([f"Sentence number {i} has enough words to split apart cleanly." for i in range(30)])
    pieces = splitter.split(text, chunk_size=64, chunk_overlap=8)
    assert len(pieces) > 1
    assert " ".join(pieces).replace(" ", "").startswith("Sentence")


def test_build_splitter_rejects_unknown() -> None:
    settings = Settings(rag_splitter="bogus")
    with pytest.raises(ConfigurationError):
        build_splitter(settings)


# --- chunk manager ------------------------------------------------------------


def test_chunk_manager_preserves_page_and_heading() -> None:
    manager = ChunkManager(
        splitter=LangChainTextSplitter(),
        chunk_size=120,
        chunk_overlap=20,
    )
    chunks = manager.chunk_document(
        [
            type(
                "Loaded",
                (),
                {
                    "page_content": ("Paragraph text with enough words. " * 20),
                    "page": 3,
                    "heading": "Methods",
                    "source": "pdf",
                    "metadata": {"source": "pdf"},
                },
            )
        ]
    )
    assert len(chunks) > 1
    assert chunks[0].page == 3
    assert chunks[0].heading == "Methods"
    assert [chunk.index for chunk in chunks] == list(range(len(chunks)))


# --- embeddings ---------------------------------------------------------------


def test_hashing_embedder_is_deterministic() -> None:
    embedder = HashingEmbedder(dims=32)
    first = embedder.embed(["quantum entanglement"])
    second = embedder.embed(["quantum entanglement"])
    assert first == second
    assert len(first[0]) == 32


def test_hashing_embedder_similarity_tracks_overlap() -> None:
    embedder = HashingEmbedder(dims=64)
    base, same, other = embedder.embed(["alpha beta gamma", "alpha beta gamma", "epsilon zeta eta"])
    assert sum(a * b for a, b in zip(base, same)) > sum(a * b for a, b in zip(base, other))


def test_build_embedder_rejects_unknown() -> None:
    with pytest.raises(ConfigurationError):
        build_embedder(Settings(embedding_provider="bogus"))


# --- vector store --------------------------------------------------------------


def test_in_memory_vector_store_search_and_delete() -> None:
    store = InMemoryVectorStore(dims=4)
    store.add(
        [[1, 0, 0, 0], [0, 1, 0, 0], [0, 0, 1, 0]],
        ["d1:c1", "d1:c2", "d1:c3"],
    )
    hits = store.search([1, 0, 0, 0], top_k=2)
    assert [hit.id for hit in hits] == ["d1:c1", "d1:c2"]
    store.delete_by_prefix("d1:")
    assert store.search([1, 0, 0, 0], top_k=3) == []


def test_faiss_vector_store_roundtrip_and_persist(tmp_path) -> None:
    index_path = tmp_path / "idx" / "index.faiss"
    store = FAISSVectorStore(dims=4, index_path=index_path)
    store.add(
        [[1, 0, 0, 0], [0, 1, 0, 0], [0, 0, 1, 0]],
        ["d1:c1", "d1:c2", "d1:c3"],
    )
    store.save()

    reloaded = FAISSVectorStore(dims=4, index_path=index_path)
    hits = reloaded.search([1, 0, 0, 0], top_k=3)
    assert hits[0].id == "d1:c1"
    assert hits[0].score > 0.99
    reloaded.delete_by_prefix("d1:")
    assert reloaded.search([1, 0, 0, 0], top_k=3) == []


def test_build_vector_store_rejects_unknown() -> None:
    with pytest.raises(ConfigurationError):
        build_vector_store(Settings(vector_store_backend="bogus"))


# --- BM25 ----------------------------------------------------------------------


def test_bm25_ranks_term_overlap() -> None:
    bm25 = BM25()
    bm25.fit(["quantum entanglement correlations", "cooking pasta recipes", "more quantum details here"])
    scores = bm25.score("quantum entanglement")
    assert scores[0] == 1.0
    assert scores[2] > scores[1]


# --- query rewrite --------------------------------------------------------------


def test_llm_query_rewriter_uses_provider_and_fails_open() -> None:
    provider = FakeLLMProvider(responses=["rewritten: best quantum phrasing"])
    rewriter = LLMQueryRewriter(provider=provider, instruction="rewrite please")
    assert rewriter.rewrite("quantum stuff") == "rewritten: best quantum phrasing"
    assert provider.calls[0][0].role == "system"

    class _Broken:
        model = "broken"

        def complete(self, messages, *, max_tokens=None, temperature=None):
            raise RuntimeError("provider down")

    assert LLMQueryRewriter(provider=_Broken(), instruction="x").rewrite("fallback query") == "fallback query"


# --- citations ------------------------------------------------------------------


def test_citation_generator_projects_hits() -> None:
    hits = [
        RetrievalHit(
            chunk_id=uuid.uuid4(),
            document_id=uuid.uuid4(),
            document_name="report.pdf",
            content="A " * 300,
            score=0.9,
            page=4,
            heading="Methods",
            fused_score=0.8,
        )
    ]
    citations = CitationGenerator().build(hits)
    assert citations[0].index == 1
    assert citations[0].document_name == "report.pdf"
    assert citations[0].page == 4
    assert citations[0].snippet == ("A " * 110).strip()


# --- retrieval service -----------------------------------------------------------


def test_retrieval_service_fuses_dense_and_lexical(rag_settings: Settings, tmp_path) -> None:
    from app.shared.database import build_engine, build_session_factory

    engine = build_engine(rag_settings.database_url)
    session_factory = build_session_factory(engine)
    Base.metadata.create_all(engine)

    store = InMemoryVectorStore(dims=32)
    embedder = HashingEmbedder(dims=32)
    service = RetrievalService(
        session_factory=session_factory,
        embedder=embedder,
        vector_store=store,
        top_k=5,
        dense_oversample=4,
        dense_weight=0.5,
        lexical_weight=0.5,
    )

    user_id = uuid.uuid4()
    doc_id = uuid.uuid4()
    from app.modules.rag.models import DocumentChunk

    with session_factory() as session:
        session.add(
            Document(
                id=doc_id,
                uploader_id=user_id,
                name="quantum.txt",
                mime="text/plain",
                size_bytes=100,
                storage_key="uploads/x/q.txt",
                status=DocumentStatus.READY.value,
                source_type="upload",
            )
        )
        session.add(DocumentChunk(document_id=doc_id, index=0, content=QUANTUM_TEXT, token_count=50))
        session.add(
            DocumentChunk(
                document_id=doc_id,
                index=1,
                content="Cooking recipes for pasta carbonara and risotto.",
                token_count=50,
            )
        )
        session.commit()
        session.flush()
        from sqlalchemy import select

        chunk_ids = session.scalars(select(DocumentChunk.id)).all()
        session.commit()

    chunk_a, chunk_b = chunk_ids
    vectors = embedder.embed([QUANTUM_TEXT, "Cooking recipes for pasta carbonara and risotto."])
    store.add(vectors, [f"{doc_id}:{chunk_a}", f"{doc_id}:{chunk_b}"])

    result = service.retrieve(user_id=user_id, query="quantum entanglement", top_k=2)
    assert len(result.hits) == 2
    assert result.hits[0].chunk_id == chunk_a
    assert result.hits[0].fused_score >= result.hits[1].fused_score
    assert result.citations[0].index == 1
    assert result.citations[0].document_name == "quantum.txt"


# --- API flow ---------------------------------------------------------------------


def test_rag_requires_auth(client: TestClient) -> None:
    assert client.post("/api/v1/rag/documents", json={"upload_id": str(uuid.uuid4())}).status_code == 401
    assert client.get("/api/v1/rag/documents").status_code == 401
    assert client.post("/api/v1/rag/retrieve", json={"query": "x"}).status_code == 401


def test_ingest_and_retrieve_roundtrip(client: TestClient) -> None:
    token = _register(client)["access_token"]
    upload = _ready_upload(
        client, token, filename="quantum.txt", content_type="text/plain", data=QUANTUM_TEXT.encode()
    )

    ingested = _ingest(client, token, upload["upload_id"])
    assert ingested["document"]["status"] == "ready"
    assert ingested["document"]["name"] == "quantum.txt"
    assert ingested["chunk_count"] >= 1
    document_id = ingested["document"]["id"]

    detail = client.get(
        f"/api/v1/rag/documents/{document_id}", headers={"Authorization": f"Bearer {token}"}
    ).json()
    assert detail["chunk_count"] == ingested["chunk_count"]

    result = _retrieve(client, token, "quantum entanglement")
    assert len(result["hits"]) >= 1
    assert result["hits"][0]["document_id"] == document_id
    assert result["hits"][0]["document_name"] == "quantum.txt"
    assert result["citations"][0]["index"] == 1
    assert result["citations"][0]["snippet"]
    assert result["rewritten_query"] == "quantum entanglement"


def test_ingest_rejects_unready_upload(client: TestClient) -> None:
    token = _register(client)["access_token"]
    result = client.post(
        "/api/v1/uploads/presign",
        json={"filename": "a.txt", "content_type": "text/plain", "size_bytes": 10},
        headers={"Authorization": f"Bearer {token}"},
    ).json()
    response = client.post(
        "/api/v1/rag/documents",
        json={"upload_id": result["upload_id"]},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 409


def test_ingest_failure_marks_document_failed(client: TestClient) -> None:
    import zipfile

    token = _register(client)["access_token"]
    import io

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("not-a-docx.txt", "just a zip")
    fake_docx = buffer.getvalue()

    upload = _ready_upload(
        client,
        token,
        filename="broken.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        data=fake_docx,
    )
    response = client.post(
        "/api/v1/rag/documents",
        json={"upload_id": upload["upload_id"]},
        headers={"Authorization": f"Bearer {token}"},
    )
    assert response.status_code == 500

    listing = client.get("/api/v1/rag/documents", headers={"Authorization": f"Bearer {token}"}).json()
    assert listing["total"] == 1
    assert listing["items"][0]["status"] == "failed"
    assert listing["items"][0]["error"]


def test_retrieval_is_user_scoped(client: TestClient) -> None:
    alice = _register(client)["access_token"]
    bob = _register(client, email="bob@example.com")["access_token"]

    upload = _ready_upload(
        client, alice, filename="quantum.txt", content_type="text/plain", data=QUANTUM_TEXT.encode()
    )
    document_id = _ingest(client, alice, upload["upload_id"])["document"]["id"]

    assert (
        client.get(
            f"/api/v1/rag/documents/{document_id}", headers={"Authorization": f"Bearer {bob}"}
        ).status_code
        == 404
    )

    bob_result = _retrieve(client, bob, "quantum entanglement")
    assert bob_result["hits"] == []

    alice_result = _retrieve(client, alice, "quantum entanglement")
    assert len(alice_result["hits"]) >= 1


def test_delete_removes_document_and_vectors(client: TestClient) -> None:
    token = _register(client)["access_token"]
    upload = _ready_upload(
        client, token, filename="quantum.txt", content_type="text/plain", data=QUANTUM_TEXT.encode()
    )
    document_id = _ingest(client, token, upload["upload_id"])["document"]["id"]

    assert _retrieve(client, token, "quantum")["hits"]

    response = client.delete(
        f"/api/v1/rag/documents/{document_id}", headers={"Authorization": f"Bearer {token}"}
    )
    assert response.status_code == 204
    assert (
        client.get(
            f"/api/v1/rag/documents/{document_id}", headers={"Authorization": f"Bearer {token}"}
        ).status_code
        == 404
    )
    assert _retrieve(client, token, "quantum")["hits"] == []


def test_list_is_user_scoped_and_paginated(client: TestClient) -> None:
    alice = _register(client)["access_token"]
    bob = _register(client, email="bob@example.com")["access_token"]

    upload_a = _ready_upload(
        client, alice, filename="one.txt", content_type="text/plain", data=b"first document content"
    )
    _ingest(client, alice, upload_a["upload_id"])
    upload_b = _ready_upload(
        client, bob, filename="two.txt", content_type="text/plain", data=b"second document content"
    )
    _ingest(client, bob, upload_b["upload_id"])

    alice_list = client.get(
        "/api/v1/rag/documents", headers={"Authorization": f"Bearer {alice}"}
    ).json()
    bob_list = client.get("/api/v1/rag/documents", headers={"Authorization": f"Bearer {bob}"}).json()
    assert alice_list["total"] == 1
    assert bob_list["total"] == 1
    assert {item["name"] for item in alice_list["items"]} == {"one.txt"}
